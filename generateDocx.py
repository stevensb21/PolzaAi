from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
import os
from docx.oxml import OxmlElement
from PIL import Image, ImageDraw, ImageFont
from datetime import datetime

certificate_details = [
    {'name': 'ПП', 'description': '«Оказание первой помощи пострадавшим (п.4в Правил)» в объеме 16 часов'}, 
    {'name': 'БГ', 'description': '"Объединенной Программе обучения безопасным методам и приемам выполнения работ при воздействии вредных и (или) опасных производственных факторов, источников опасности, идентифицированных в рамках специальной оценки условий труда и оценки профессиональных рисков; по использованию (применению) средств индивидуальной защиты для работников рабочих профессий (п.46б, п.4г Правил)" в объеме 24 часа'}, 
    {'name': 'вР (46в)', 'description': '"Программе обучения безопасным методам и приемам выполнения работ повышенной опасности, к которым предъявляются дополнительные требования в соответствии с нормативными правовыми актами, содержащими государственные нормативные требования охраны труда для работников, непосредственно выполняющих работы повышенной опасности (п.46в Правил)" в объеме 16 часов'},
    {'name': 'Высота(рабочая, 2гр)', 'description': '«Безопасные методы и приемы выполнения работ на высоте для работников 2-ой группы» в объеме 32 часа'},
]
def generate_filename(order_data):
    """Генерирует имя файла по шаблону [дата]-[Фамилия]-[Имя первых двух удостоверений]"""
    # Получаем текущую дату
    current_date = datetime.now().strftime("%d.%m.%Y")
    
    # Извлекаем фамилию из полного имени
    full_name = order_data['employee']['full_name']
    surname = full_name.split()[0] if full_name else "Unknown"
    
    # Получаем названия первых двух удостоверений
    cert_names = []
    for cert in certificate_details[:2]:  # Берем только первые два
        cert_names.append(cert['name'])
    
    # Объединяем названия удостоверений
    cert_part = "-".join(cert_names) if cert_names else "Unknown"
    
    # Формируем имя файла
    filename = f"{current_date}-{surname}-{cert_part}.docx"
    
    # Заменяем недопустимые символы для имени файла
    invalid_chars = ['<', '>', ':', '"', '/', '\\', '|', '?', '*']
    for char in invalid_chars:
        filename = filename.replace(char, '_')
    
    return filename


def create_tetracom_document(order_data):
    global certificate_details
    # Создаем документ
    doc = Document()
    
    # Берем первую секцию документа
    section = doc.sections[0]
    
    # Настройка полей страницы
    section.top_margin = Inches(0.3)      # верхнее поле
    section.bottom_margin = Inches(1)     # нижнее поле
    section.left_margin = Inches(0.3)       # левое поле
    section.right_margin = Inches(0.3)      # правое поле
    section.header_distance = Inches(0.3) # расстояние от верхнего поля до колонтитула
    section.footer_distance = Inches(0.3) # расстояние от нижнего поля до колонтитула
    
    # Настройки для изображения в колонтитуле
    image_margin = Inches(0.1)  # Дополнительный отступ для красоты
    
    # -------------------
    # Верхний колонтитул с изображением
    # -------------------
    header = section.header
    header_paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Проверяем, есть ли изображение
    image_path = 'images/header_logo.png'  # Путь к изображению
    if os.path.exists(image_path):
        # Добавляем изображение в верхний колонтитул
        # Вычисляем ширину страницы минус отступы (с дополнительным отступом для красоты)
        page_width = section.page_width - section.left_margin - section.right_margin - image_margin
        header_paragraph.add_run().add_picture(image_path, width=page_width)
    else:
        # Если изображения нет, добавляем текст
        run = header_paragraph.add_run("""Общество с ограниченной ответственностью «Тетраком»
420194, Республика Татарстан (Татарстан), 
г. Казань, ул. Короленко, д.115, помещ.53
ОГРН 1121690081853, ИНН 1661034770, КПП 166101001
""")
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        run.bold = True
    


    doc.add_paragraph('')
    # -------------------
    # Основное содержимое документа
    # -------------------
    # Добавляем заголовок
    title = doc.add_heading("""Общество с ограниченной ответственностью «Тетраком»
420194, Республика Татарстан (Татарстан), 
г. Казань, ул. Короленко, д.115, помещ.53
ОГРН 1121690081853, ИНН 1661034770, КПП 166101001
""", 0)
    
    # Настраиваем шрифт для заголовка
    for run in title.runs:
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        run.font.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run.font.color.rgb = RGBColor(0, 0, 0)  # Черный цвет
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Устанавливаем межстрочный интервал (1.5 строки)
    title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    title.paragraph_format.line_spacing = 1
    
    # Добавляем блок с исходящим номером и датой
    # Создаем таблицу для размещения элементов в одной строке
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Левая ячейка - исходящий номер
    left_cell = table.cell(0, 0)
    left_para = left_cell.paragraphs[0]
    left_run = left_para.add_run("Исх.№ ________________")
    left_run.font.size = Pt(11)
    left_run.font.name = 'Calibri'
    left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Правая ячейка - дата
    right_cell = table.cell(0, 1)
    right_para = right_cell.paragraphs[0]
    right_run = right_para.add_run("«___» _______ 20___ г.")
    right_run.font.size = Pt(11)
    right_run.font.name = 'Calibri'
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    

    
    # Добавляем информацию о получателе (выровнено по правому краю)
    recipient_para = doc.add_paragraph()
    recipient_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    recipient_run1 = recipient_para.add_run("Директору АНО ДПО")
    recipient_run1.font.size = Pt(11)
    recipient_run1.font.name = 'Calibri'
    
    recipient_para.add_run('\n')
    
    recipient_run2 = recipient_para.add_run("«РУЦПК»")
    recipient_run2.font.size = Pt(11)
    recipient_run2.font.name = 'Calibri'
    
    recipient_para.add_run('\n')
    
    recipient_run3 = recipient_para.add_run("Грачевой А.А.")
    recipient_run3.font.size = Pt(11)
    recipient_run3.font.name = 'Calibri'
    
    recipient_para.add_run('\n')
    
    recipient_run4 = recipient_para.add_run("email: ruts_pk@mail.ru")
    recipient_run4.font.size = Pt(11)
    recipient_run4.font.name = 'Calibri'
    

    p1 = doc.add_paragraph('Уважаемая Асия Амировна!')    
    # Настраиваем шрифт для параграфа
    for run in p1.runs:
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0, 0, 0)  # Черный цвет
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Устанавливаем межстрочный интервал
    p1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p1.paragraph_format.line_spacing = 1
    

    p2 = doc.add_paragraph('Прошу Вас провести обучение по программам:')    
    # Настраиваем шрифт для параграфа
    for run in p2.runs:
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0, 0, 0)  # Черный цвет
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT



    # Добавляем информацию о продлении
    for index, certifcate in enumerate(certificate_details, 1):
        cert = doc.add_paragraph(f'{index}. ' + certifcate['name'] + ' - ' + certifcate['description'])
        for run in cert.runs:
            run.font.size = Pt(11)
            run.font.name = 'Calibri'
            run.font.color.rgb = RGBColor(0, 0, 0)  # Черный цвет
        cert.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Добавляем отступ слева как у параграфа
        cert.paragraph_format.left_indent = Inches(0.3)
        # Убираем отступы до и после
        cert.paragraph_format.space_before = Pt(0)
        cert.paragraph_format.space_after = Pt(0)
        # Настраиваем параграф так, чтобы он не разрывался с таблицей
        cert.paragraph_format.keep_with_next = True
    doc.add_paragraph('Cледующим работникам:')    
        
        
    # Создаем таблицу с данными сотрудника
    table = doc.add_table(rows=2, cols=6)  # 2 строки: заголовок + данные
    table.style = 'Table Grid'
    
    # Запрещаем разрыв между строками таблицы
    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        cantSplit = OxmlElement('w:cantSplit')
        trPr.append(cantSplit)
    
    
    # Заголовки таблицы
    header_cells = table.rows[0].cells
    headers = ['№ \n', 'ФИО \n', 'Год рождения \n', 'Должность \n', 'СНИЛС \n', 'ИНН \n']
    
    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.text = header
        # Настраиваем шрифт для заголовков
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(11)
                run.font.name = 'Calibri'
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Данные сотрудника
    data_cells = table.rows[1].cells
    employee = order_data['employee']
    
    # Форматируем дату рождения в полном формате дд.мм.гггг
    birth_date = employee['birth_date']
    if birth_date and birth_date != "null":
        try:
            # Если дата в формате ISO (1974-12-20T00:00:00.000000Z)
            if "T" in birth_date:
                from datetime import datetime
                date_obj = datetime.fromisoformat(birth_date.replace("Z", "+00:00"))
                birth_date_formatted = date_obj.strftime("%d.%m.%Y")
            # Если дата уже в формате дд.мм.гггг
            elif "." in birth_date and len(birth_date.split(".")) == 3:
                birth_date_formatted = birth_date
            else:
                birth_date_formatted = birth_date
        except Exception as e:
            birth_date_formatted = birth_date
    else:
        birth_date_formatted = "не указана"
    
    data = [
        '1'+'\n',  # №
        employee['full_name'],  # ФИО
        birth_date_formatted,  # Полная дата рождения в формате дд.мм.гггг
        employee['position'],  # Должность
        employee['snils'],  # СНИЛС
        employee['inn']  # ИНН
    ]
    
    for i, value in enumerate(data):
        cell = data_cells[i]
        cell.text = value
        # Настраиваем шрифт для данных
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(11)
                run.font.name = 'Calibri'
                run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Добавляем отступ после таблицы
    doc.add_paragraph('')
    
    # Добавляем "Директор"
    director_para = doc.add_paragraph('Директор')
    for run in director_para.runs:
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0, 0, 0)
        
    director_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Добавляем отступ слева как у параграфа
    director_para.paragraph_format.left_indent = Inches(0.3)
    # Убираем интервал после параграфа
    director_para.paragraph_format.space_after = Pt(0)
    
    # Создаем таблицу для "ООО «ТЕТРАКОМ»" и "Д.И. Дворяшов"
    signature_table = doc.add_table(rows=1, cols=2)
    
    
    # Левая ячейка - ООО «ТЕТРАКОМ»
    company_cell = signature_table.cell(0, 0)
    company_para = company_cell.paragraphs[0]
    company_run = company_para.add_run('ООО «ТЕТРАКОМ»')
    company_run.font.size = Pt(11)
    company_run.font.name = 'Calibri'
    company_run.font.color.rgb = RGBColor(0, 0, 0)
    company_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Добавляем отступ слева как у параграфа
    company_para.paragraph_format.left_indent = Inches(0.3)
    
    # Правая ячейка - Д.И. Дворяшов
    name_cell = signature_table.cell(0, 1)
    name_para = name_cell.paragraphs[0]
    name_run = name_para.add_run('Д.И. Дворяшов')
    name_run.font.size = Pt(11)
    name_run.font.name = 'Calibri'
    name_run.font.color.rgb = RGBColor(0, 0, 0)
    name_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # Добавляем отступ справа как у параграфа
    name_para.paragraph_format.right_indent = Inches(0.3)
    

    
    # -------------------
    # Нижний колонтитул с изображением
    # -------------------
    footer = section.footer
    footer_paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Проверяем, есть ли изображение
    if os.path.exists(image_path):
        # Добавляем изображение в нижний колонтитул
        # Вычисляем ширину страницы минус отступы (с дополнительным отступом для красоты)
        page_width = section.page_width - section.left_margin - section.right_margin - image_margin
        footer_paragraph.add_run().add_picture(image_path, width=page_width)
    else:
        # Если изображения нет, добавляем текст
        run = footer_paragraph.add_run("Общество с ограниченной ответственностью «Тетраком»")
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        run.bold = True
    
    # -------------------
    # Сохраняем документ
    # -------------------
    filename = generate_filename(order_data)
    doc.save(filename)
    print(f"Документ успешно создан: {filename}")
    return filename

if __name__ == "__main__":
   
    order_data = {'type': 'readyorder', 'employee': {'id': 92, 'full_name':
  'Ахмедшин Альберт Рашитович', 'snils': '079-205-590-83', 'inn': '4949495566', 'position':
  'Монтажник', 'birth_date': '20.12.1974', 'phone': '47484885858859', 'photo':
  'https://api.telegram.org/file/bot7934916395:AAFKUivVycRo_sf2azOv4oJUVaslRJ1j1YU/photos/file_6.jpg'}}
    
    create_tetracom_document(order_data)
